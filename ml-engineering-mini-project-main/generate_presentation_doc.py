"""One-off script to generate ETA_Presentation_Data_Pipeline_Dashboard.docx -
a presentation-focused document: raw data sample, what was modified, how the
pipeline is built/triggered, FastAPI endpoints, the feedback loop, and the
dashboards with embedded screenshots and the exact datapoints each is built
from. Not part of the ML pipeline; run manually to regenerate.
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "ETA_Presentation_Data_Pipeline_Dashboard.docx"
REPORTS = ROOT / "reports"

BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x8E, 0x3E)
RED = RGBColor(0xC0, 0x39, 0x2B)


def add_heading(doc, text, level=1, color=BLUE):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_image(doc, path, width_in=6.0, caption=None):
    doc.add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY


RAW_HEAD = """     trip_id  pickup_lat  pickup_lon  dropoff_lat  dropoff_lon  trip_distance_km  hour  weekday weather traffic_level  trip_duration_minutes            created_at
0  TRIP00528   40.801834  -73.834051    40.689163   -73.872286             12.67    20        3  cloudy        medium                  41.19  2026-08-24T20:15:00Z
1  TRIP00095   40.782317  -73.825760    40.732624   -74.016020              7.24     4        4  cloudy          high                  32.26  2026-08-11T04:15:00Z
2  TRIP00044   40.802207  -73.844582    40.699836   -73.940837              2.50    22        5  cloudy          high                  32.62  2026-08-16T22:15:00Z
3  TRIP00590   40.837177  -73.892724    40.600579   -73.896548             19.27     2        2   clear        medium                  53.68  2026-08-02T02:15:00Z
4  TRIP00259   40.616143  -73.999973    40.780750   -73.897251             27.06    11        4    snow          high                  93.48  2026-08-07T11:15:00Z"""

RAW_DESCRIBE = """       pickup_lat  pickup_lon  dropoff_lat  dropoff_lon  trip_distance_km        hour     weekday  trip_duration_minutes
count  602.000000  603.000000   603.000000   603.000000        603.000000  603.000000  603.000000             603.000000
mean    40.724227  -73.927626    40.722521   -73.921609         14.961509   11.502488    3.149254              51.023516
std      0.072193    0.071751     0.073634     0.073497          7.941306    6.785605    1.985231              19.599961
min     40.600104  -74.049658    40.600075   -74.049966         -2.000000    0.000000    0.000000               6.410000
25%     40.662052  -73.987605    40.663460   -73.983694          8.205000    6.000000    1.500000              36.615000
50%     40.723809  -73.929927    40.716938   -73.920260         15.630000   12.000000    3.000000              52.050000
75%     40.788893  -73.867481    40.784390   -73.859137         21.975000   17.000000    5.000000              65.890000
max     40.849875  -73.800226    40.849976   -73.800000         27.920000   23.000000    6.000000              93.480000"""

CLEAN_DESCRIBE = """       pickup_lat  pickup_lon  dropoff_lat  dropoff_lon  trip_distance_km        hour     weekday  trip_duration_minutes
count  600.000000  600.000000   600.000000   600.000000        600.000000  600.000000  600.000000             600.000000
mean    40.724308  -73.927764    40.722634   -73.922217         15.022983   11.520000    3.160000              51.178633
std      0.072299    0.071904     0.073801     0.073174          7.909791    6.798036    1.984346              19.525318
min     40.600104  -74.049658    40.600075   -74.049966          1.000000    0.000000    0.000000               6.410000
25%     40.661918  -73.987831    40.662766   -73.983978          8.242500    6.000000    2.000000              37.187500
50%     40.725274  -73.930051    40.717405   -73.920764         15.685000   12.000000    3.000000              52.245000
75%     40.789088  -73.866628    40.784728   -73.860911         21.995000   17.000000    5.000000              65.970000
max     40.849875  -73.800226    40.849976   -73.800142         27.920000   23.000000    6.000000              93.480000"""


def build():
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading(level=0)
    run = title.add_run("ETA Prediction Project")
    run.font.color.rgb = BLUE
    add_para(doc, "Data, Pipeline, API, Feedback Loop, and Dashboards - Presentation Notes", italic=True, color=GRAY)
    doc.add_paragraph()

    # ---------------- 1. Raw data sample ----------------
    add_heading(doc, "1. Raw Data Sample (data/raw/trips.csv)")
    add_para(doc, "pd.read_csv('data/raw/trips.csv').head()  ->  603 rows total (600 valid + 3 intentionally invalid)", bold=True)
    add_code_block(doc, RAW_HEAD)
    add_para(doc, "pd.read_csv('data/raw/trips.csv').describe()", bold=True)
    add_code_block(doc, RAW_DESCRIBE)
    add_para(doc, "Notice the 3 injected bad rows show up here: pickup_lat count is 602 (one missing value), "
                  "trip_distance_km min is -2.0 (an impossible negative distance), and one row has an "
                  "unparseable created_at timestamp (not visible in describe() since it's a string column). "
                  "These are the exact records validation is designed to catch.", italic=True, color=GRAY)

    # ---------------- 2. What we modified ----------------
    add_heading(doc, "2. Validation - What We Modified (Raw -> Clean)")
    add_para(doc, "pipeline/validate.py reads data/raw/trips.csv and checks every row against these rules:")
    add_bullets(doc, [
        "Schema: column headers must exactly match the expected 12 fields",
        "GPS bounds: pickup/dropoff latitude in [-90, 90], longitude in [-180, 180]",
        "Distance: trip_distance_km must be > 0 and <= 200",
        "Time: hour in [0, 23], weekday in [0, 6]",
        "Categorical: weather in {clear, rain, snow, cloudy}; traffic_level in {low, medium, high}",
        "Timestamp: created_at must parse as ISO-8601",
        "Duration: trip_duration_minutes must be > 0",
    ])
    add_table(doc, ["", "Raw (data/raw/trips.csv)", "Clean (data/processed/trips_clean.csv)"], [
        ["Row count", "603", "600"],
        ["Rejected rows", "-", "3 (written to data/processed/rejected_rows.jsonl with their error reasons)"],
        ["pickup_lat count", "602 (1 missing)", "600 (no missing values)"],
        ["trip_distance_km min", "-2.00 (invalid)", "1.00 (valid)"],
        ["created_at", "1 row unparseable", "all rows parse as ISO-8601"],
    ], col_widths=[1.6, 2.7, 2.7])
    add_para(doc, "pd.read_csv('data/processed/trips_clean.csv').describe()", bold=True)
    add_code_block(doc, CLEAN_DESCRIBE)
    add_para(doc, "The clean dataset is then hashed and versioned: data/manifest.json stores the SHA-256 hash, "
                  "row counts, and seed; data/VERSION stores the dataset version string (dataset-a-v1.0.0) so "
                  "every training run can be traced back to the exact data it used.")

    # ---------------- 3. Feature set ----------------
    add_heading(doc, "3. From Clean Data to Model Features")
    add_table(doc, ["Type", "Columns", "Transformation before reaching the model"], [
        ["Numeric", "pickup_lat, pickup_lon, dropoff_lat, dropoff_lon, trip_distance_km, hour, weekday", "StandardScaler (fit on training split only)"],
        ["Categorical", "weather, traffic_level", "OneHotEncoder(handle_unknown='ignore')"],
        ["Target", "trip_duration_minutes", "Not transformed - this is what the model predicts"],
        ["Dropped", "trip_id, created_at", "Identifiers/metadata only, not predictive features"],
    ], col_widths=[1.1, 3.4, 2.5])
    add_para(doc, "Both transforms are combined in one scikit-learn ColumnTransformer (pipeline/features.py + "
                  "pipeline/train.py: make_preprocessor()), and that fitted object is saved together with the "
                  "model inside artifacts/model.joblib - so the API reuses the exact same fitted scaler/encoder "
                  "at prediction time.")
    add_para(doc, "Known gap: hour and weekday are passed through as raw numbers rather than being converted into "
                  "explicit derived flags such as is_weekend or is_rush_hour.", italic=True, color=RED)

    # ---------------- 4. Pipeline ----------------
    add_heading(doc, "4. How the Pipeline Is Built and Triggered")
    add_para(doc, "Everything lives in pipeline/train.py as a single callable main() function:")
    add_bullets(doc, [
        "1. generate_data.main() - (re)generates the deterministic raw dataset",
        "2. validate.validate() - validates it, produces the clean CSV + rejected rows",
        "3. Load clean rows, split 80/20 train/test (seed=412, from params.yaml)",
        "4. Build the shared preprocessor (ColumnTransformer)",
        "5. Fit Linear Regression and Gradient Boosting, each wrapped with the preprocessor in one sklearn Pipeline",
        "6. Evaluate both on the held-out test set (MAE, RMSE, R2)",
        "7. Log every run's params/metrics/tags to MLflow (experiment: 'delivery-eta')",
        "8. Select the model with the lowest test RMSE",
        "9. Persist the winning fitted Pipeline + metadata to artifacts/model.joblib",
        "10. Write reports/model_comparison.md, reports/training_report.json, experiments/runs.jsonl",
    ])
    add_para(doc, "How it's triggered:", bold=True)
    add_bullets(doc, [
        "Manually: python -m pipeline.train (runs locally against the current Python environment)",
        "Via Docker (what run.bat does): docker run --rm -v \"%DIR%:/app\" -w /app ml-engineering-mini-project python -m pipeline.train",
        "run.bat automates the full sequence: build image -> start MLflow -> train -> run monitoring -> run tests -> rebuild image with the trained model baked in -> start the API container",
        "There is no scheduled/automatic trigger yet (e.g., cron or a watcher on the drift reports) - training is currently triggered manually or via run.bat",
    ])

    # ---------------- 5. FastAPI endpoints ----------------
    add_heading(doc, "5. FastAPI Endpoints Used (api/server.py)")
    add_table(doc, ["Endpoint", "Method", "Purpose", "Request", "Response"], [
        ["/predict", "POST", "Return a predicted ETA for a trip",
         "TripRequest: pickup_lat, pickup_lon, dropoff_lat, dropoff_lon, trip_distance_km, hour, weekday, weather, traffic_level (Pydantic-validated ranges/enums)",
         "TripResponse: predicted_eta_minutes, model_version, latency_ms, request_id"],
        ["/health", "GET", "Liveness check", "none", "{status: 'ok', model_version}"],
        ["/feedback", "POST", "Submit the real outcome for a previous prediction",
         "FeedbackRequest: request_id, actual_eta_minutes (0-500), labelled_at_utc (optional)",
         "{status: 'recorded', request_id, actual_eta_minutes, labelled_at_utc}"],
    ], col_widths=[0.9, 0.7, 1.6, 2.5, 1.8])
    add_para(doc, "Invalid input never reaches the model: Pydantic validation errors are caught by a custom "
                  "exception handler and return HTTP 422 with structured error details. Every successful "
                  "/predict call is logged to logs/predictions.jsonl (timestamp, request_id, input hash, "
                  "prediction, model version, latency); every /feedback call is logged to logs/feedback.jsonl.")

    # ---------------- 6. Feedback loop ----------------
    add_heading(doc, "6. How the Feedback Loop Works")
    add_bullets(doc, [
        "client/live_client.py sends simulated real-time trip requests to POST /predict and logs the request+response to logs/client_requests.jsonl/.txt",
        "client/feedback_client.py later simulates the real-world outcome for each logged request (using the same physics-based duration formula as the data generator, plus noise; a --surge flag adds a festival/rush-hour delay bonus on demand) and submits it via POST /feedback",
        "pipeline/accuracy_report.py joins logs/predictions.jsonl with logs/feedback.jsonl by request_id, computing error = actual - predicted for every matched pair",
        "It reports overall MAE/RMSE, plus MAE/RMSE split into an older half and a newer half of matched pairs - a rising RMSE in the newer window sets accuracy_drift_triggered = true once it crosses 8.0 minutes",
        "This produces reports/accuracy_drift.md/.json, which both feed the observability dashboard's feedback-loop section",
    ])
    add_para(doc, "This is the real 'predicted vs. actual' signal - distinct from feature-level PSI/TVD drift, "
                  "which only checks whether inputs look statistically different from training data.", italic=True, color=GRAY)

    # ---------------- 7. Dashboards with screenshots ----------------
    add_heading(doc, "7. Dashboards - Datapoints and Screenshots")

    add_heading(doc, "7.1 Health / Monitoring Dashboard (reports/dashboard.html)", level=2)
    add_para(doc, "Built by pipeline/build_dashboard.py. Datapoints used, all read from logs/client_requests.jsonl "
                  "(one JSON record per simulated client request):")
    add_bullets(doc, [
        "timestamp_utc - x-axis for both time-series charts",
        "outcome (ok / http_error / crash) - determines point/bar color (green / orange / red)",
        "status_code - plotted directly in the status-code timeline (0 represents a crash/no response)",
        "latency_ms - plotted in the latency-over-time chart",
    ])
    add_para(doc, "Computed summary: total requests, uptime % = ok / total, counts of 200 OK / HTTP errors / crashes.")
    if (REPORTS / "_screenshot_dashboard.png").exists():
        add_image(doc, REPORTS / "_screenshot_dashboard.png", width_in=6.0,
                   caption="Figure 1: Health dashboard - 202 requests, 99.01% uptime. The latency spike (~4,000 ms) "
                           "and the single orange bar are the deliberately injected crash and invalid-payload test cases.")

    add_heading(doc, "7.2 Observability Dashboard, Part A - Input Feature Drift", level=2)
    add_para(doc, "Built by pipeline/build_observability_dashboard.py from reports/drift_table.json (produced by "
                  "pipeline/drift_report.py, which compares data/processed/trips_clean.csv against live requests "
                  "logged in logs/client_requests.jsonl). Datapoints used per feature:")
    add_bullets(doc, [
        "baseline_mean / live_mean, baseline_std / live_std - summary statistics compared side by side",
        "psi - Population Stability Index for numeric features (pickup/dropoff lat-lon, distance, hour, weekday)",
        "total_variation_distance - category-distribution distance for weather and traffic_level",
        "mean_shift_z - how many baseline standard deviations the live mean has moved",
        "drift (boolean) - true if psi >= 0.2 (numeric) or TVD >= 0.15 (categorical)",
    ])
    if (REPORTS / "_obs_section_a.png").exists():
        add_image(doc, REPORTS / "_obs_section_a.png", width_in=6.0,
                   caption="Figure 2: Overview cards + PSI per numeric feature - all 7 numeric features are under the 0.2 threshold (no drift).")
    if (REPORTS / "_obs_section_b.png").exists():
        add_image(doc, REPORTS / "_obs_section_b.png", width_in=6.0,
                   caption="Figure 3: TVD per categorical feature + detail tables - 'weather' is flagged (0.2567 > 0.15) because the simulated client over-samples 'clear' weather versus training's uniform distribution.")

    add_heading(doc, "7.3 Observability Dashboard, Part B - Predicted vs. Actual Accuracy Drift", level=2)
    add_para(doc, "Same dashboard, built from reports/accuracy_drift.json (produced by pipeline/accuracy_report.py "
                  "joining logs/predictions.jsonl with logs/feedback.jsonl). Datapoints used:")
    add_bullets(doc, [
        "predicted_eta_minutes vs. actual_eta_minutes per matched request_id, plotted over time as two lines",
        "abs_error = |actual - predicted| per request, plotted as bars against the RMSE drift threshold (8.0 minutes)",
        "matched_pairs, overall MAE/RMSE, and the older-half vs. newer-half RMSE comparison feeding accuracy_drift_triggered",
    ])
    if (REPORTS / "_obs_section_c.png").exists():
        add_image(doc, REPORTS / "_obs_section_c.png", width_in=6.0,
                   caption="Figure 4: Predicted vs. actual ETA over time and absolute error per request - 120 matched pairs, RMSE 13.56 (above the 8.0 threshold), accuracy_drift_triggered = YES, driven by the simulated surge scenario.")

    doc.save(OUT_PATH)
    print(f"Saved to {OUT_PATH}")


if __name__ == "__main__":
    build()
