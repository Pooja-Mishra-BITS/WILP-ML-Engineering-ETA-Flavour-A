"""One-off script to generate ETA_Project_Explanation.docx summarizing the
project's architecture, folder structure, and week-by-week (M2-M5) coverage
for the presentation. Not part of the pipeline; run manually if the doc needs
regenerating.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "ETA_Project_Explanation.docx"

GREEN = RGBColor(0x1E, 0x8E, 0x3E)
RED = RGBColor(0xC0, 0x39, 0x2B)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level=1, color=BLUE):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading(level=0)
    run = title.add_run("Delivery / Ride ETA Prediction")
    run.font.color.rgb = BLUE
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("End-to-end ML Engineering Project - Architecture, Folder Structure, and Weekly Coverage")
    r.italic = True
    r.font.color.rgb = GRAY
    doc.add_paragraph()

    # ---------------- Overview ----------------
    add_heading(doc, "1. Project Overview")
    add_para(doc,
        "This project predicts delivery/ride ETA (estimated time of arrival) using trip distance, pickup/dropoff "
        "location, time of day, weather, and traffic level. It covers the complete ML engineering lifecycle:")
    add_para(doc, "Raw Data -> Validation -> Feature Set -> Training -> MLflow Tracking -> Model Selection -> "
                  "REST API -> Client Traffic Simulation -> Logging -> Drift & Accuracy Monitoring -> Dashboards",
             bold=True, color=BLUE)
    add_para(doc, "The dataset is synthetic (deterministic, seed 412) and built purely to demonstrate the "
                  "engineering workflow end-to-end, not as a substitute for real fleet telemetry.")

    # ---------------- Architecture diagram (text) ----------------
    add_heading(doc, "2. High-Level Architecture")
    add_para(doc,
        "data/raw/trips.csv  -->  pipeline/validate.py  -->  data/processed/trips_clean.csv\n"
        "  -->  pipeline/train.py (Linear Regression vs Gradient Boosting, tracked in MLflow)\n"
        "  -->  artifacts/model.joblib (best model selected by lowest RMSE)\n"
        "  -->  api/server.py  (FastAPI: /predict, /health, /feedback)\n"
        "  -->  client/live_client.py (simulated real-time traffic) --> logs/predictions.jsonl, "
        "logs/client_requests.txt/.jsonl\n"
        "  -->  client/feedback_client.py (simulated ground-truth outcomes) --> logs/feedback.jsonl\n"
        "  -->  pipeline/drift_report.py (input feature drift, PSI/TVD)\n"
        "  -->  pipeline/accuracy_report.py (predicted-vs-actual accuracy drift)\n"
        "  -->  pipeline/build_dashboard.py + build_observability_dashboard.py --> reports/*.html")

    # ---------------- Folder structure ----------------
    add_heading(doc, "3. Folder Structure")
    add_para(doc, "Each top-level folder maps to one stage of the ML lifecycle:")
    add_table(doc, ["Folder / File", "Purpose"], [
        ["data/", "Raw, validated, and processed trip data; dataset version (VERSION) and hash manifest (manifest.json)"],
        ["pipeline/", "All pipeline logic: data generation, validation, feature list, training, monitoring, drift & accuracy reports, dashboards"],
        ["artifacts/", "The single deployed model file (model.joblib) - winner of the model comparison, with its preprocessing pipeline bundled in"],
        ["mlruns/", "MLflow's experiment tracking store - every training run's parameters, metrics, and tags"],
        ["experiments/", "Human-readable run results (runs.jsonl) plus a copy of the selected model"],
        ["api/", "The FastAPI service - /predict (returns ETA), /health (liveness), /feedback (submit actual outcomes)"],
        ["client/", "live_client.py simulates real users querying the API; feedback_client.py simulates real-world outcomes coming back later"],
        ["logs/", "Everything observed in production: predictions.jsonl, feedback.jsonl, client_requests.txt/.jsonl"],
        ["reports/", "All generated reports and dashboards - what you actually present in the demo"],
        ["tests/", "Automated unit tests (API contract, validation errors, health check)"],
        ["submission/evidence/", "Written evidence/checklists mapping deliverables to grading criteria"],
        ["run.bat", "One-command setup: build image, start MLflow, train, run monitoring, run tests, start API"],
        ["run_client.bat", "One-command demo: simulate traffic, submit feedback, refresh drift/accuracy reports and both dashboards"],
        ["Dockerfile / requirements.txt / params.yaml", "Container definition, pinned dependencies, and training configuration (seed, test split, model hyperparameters)"],
    ], col_widths=[1.8, 4.7])

    # ---------------- Week by week ----------------
    add_heading(doc, "4. Week-by-Week Coverage (M2-M5)")

    # Week 1
    add_heading(doc, "Week 1 - M2: Ingest, Validate, Engineer Features, Version Dataset", level=2)
    add_para(doc, "Goal: Ingest historical trip data; validate schema (missing GPS pings, invalid timestamps); "
                  "engineer features (hour-of-day, weekday/weekend, distance, weather); version the dataset.")
    add_bullets(doc, [
        "Ingestion: pipeline/generate_data.py produces 603 deterministic trips (600 valid + 3 intentionally invalid) into data/raw/trips.csv",
        "Validation: pipeline/validate.py checks schema match, GPS bounds, distance range, hour/weekday range, weather/traffic enums, and ISO timestamp parsing -> data/processed/trips_clean.csv + rejected_rows.jsonl",
        "Feature set: pipeline/features.py defines numeric features (pickup/dropoff lat-lon, distance, hour, weekday) and categorical features (weather, traffic level)",
        "Dataset versioning: data/VERSION and data/manifest.json (seed, row counts, SHA-256 hash)",
    ])
    add_para(doc, "Status: Mostly complete. Note: feature engineering is currently limited to raw columns "
                  "(hour, weekday used directly) rather than fully derived features such as an explicit "
                  "is_weekend or is_rush_hour flag - a good talking point for future improvement.",
             italic=True, color=RED)

    # Week 2
    add_heading(doc, "Week 2 - M3: Train & Compare Models, Track Experiments", level=2)
    add_para(doc, "Goal: Train and compare models (linear regression vs. gradient boosting); track experiments "
                  "and hyperparameters.")
    add_bullets(doc, [
        "pipeline/train.py trains Linear Regression and Gradient Boosting side by side using an identical preprocessing pipeline (StandardScaler + OneHotEncoder)",
        "Every run is logged to MLflow (params, metrics: MAE/RMSE/R2, tags: dataset version, feature representation)",
        "The model with the lowest held-out RMSE is automatically selected",
        "Results: reports/model_comparison.md and experiments/runs.jsonl",
    ])
    add_para(doc, "Status: Fully complete.", italic=True, color=GREEN)

    # Week 3
    add_heading(doc, "Week 3 - M4: Package and Serve the Best Model via REST API", level=2)
    add_para(doc, "Goal: Package the best model; serve via a REST API that accepts trip details and returns "
                  "predicted ETA.")
    add_bullets(doc, [
        "Best model + preprocessing pipeline packaged into artifacts/model.joblib with metadata (model_version, dataset_version, seed, features)",
        "api/server.py - FastAPI service exposing POST /predict (validated with Pydantic, returns predicted_eta_minutes, model_version, latency_ms, request_id) and GET /health",
        "Deployed as a Docker container, verified live and reachable on http://localhost:8000",
    ])
    add_para(doc, "Status: Fully complete.", italic=True, color=GREEN)

    # Week 4
    add_heading(doc, "Week 4 - M5: Log Predictions vs. Actual, Simulate Drift, Monitor, Retraining Trigger", level=2)
    add_para(doc, "Goal: Log predictions vs. actual times; simulate drift (e.g., festival/rush-hour surge); "
                  "set up monitoring and a retraining trigger.")
    add_bullets(doc, [
        "Prediction logging: logs/predictions.jsonl (API-side) and logs/client_requests.txt/.jsonl (client-side, human + machine readable)",
        "Feedback loop (predicted vs. actual): client/feedback_client.py simulates real-world outcomes for logged requests and submits them via POST /feedback -> logs/feedback.jsonl",
        "Accuracy drift: pipeline/accuracy_report.py joins predictions with feedback by request_id, computes MAE/RMSE overall and in older-vs-newer time windows, and raises accuracy_drift_triggered when newer-window RMSE crosses a threshold -> reports/accuracy_drift.md",
        "Simulated drift: pipeline/monitor.py generates a scripted festival/rush-hour surge batch (heavier traffic, adverse weather, extra congestion delay) and compares it to a training baseline using PSI",
        "Live input drift: pipeline/drift_report.py computes PSI (numeric features) and Total Variation Distance (categorical features) between real live client traffic and the training baseline -> reports/drift_table.md",
        "Monitoring dashboards: reports/dashboard.html (uptime %, status codes, crash timeline, latency) and reports/observability_dashboard.html (PSI/TVD per feature plus the predicted-vs-actual accuracy drift charts)",
    ])
    add_para(doc, "Status: Mostly complete. Note: retraining is currently a computed decision signal "
                  "(retraining_triggered / accuracy_drift_triggered flags in the reports) rather than an "
                  "automated action - nothing yet watches these flags and invokes pipeline.train automatically. "
                  "Good next step to mention.", italic=True, color=RED)

    # ---------------- Feature engineering ----------------
    add_heading(doc, "5. Feature Engineering - What We Actually Did")
    add_para(doc, "The model consumes 9 raw trip attributes, split into two groups that are transformed "
                  "differently before reaching the model:")
    add_table(doc, ["Type", "Features", "Transformation Applied"], [
        ["Numeric", "pickup_lat, pickup_lon, dropoff_lat, dropoff_lon, trip_distance_km, hour, weekday",
         "StandardScaler (zero mean, unit variance) - puts all numeric columns on a comparable scale so no single feature dominates due to magnitude"],
        ["Categorical", "weather, traffic_level",
         "OneHotEncoder (handle_unknown='ignore') - converts each category into a binary column so the model can use it numerically"],
    ], col_widths=[1.2, 3.0, 3.3])
    add_para(doc, "Both transformations are combined with a scikit-learn ColumnTransformer, then wrapped together "
                  "with the regressor inside a single sklearn Pipeline object (pipeline/train.py -> "
                  "make_preprocessor()). This means scaling and encoding are learned only on the training split "
                  "and then reused identically at prediction time - the exact same fitted preprocessor is saved "
                  "inside artifacts/model.joblib, so the API never has to re-derive it.")
    add_para(doc, "Known limitation (worth stating openly in the presentation): hour and weekday are used as raw "
                  "numbers rather than being turned into explicit engineered signals such as is_weekend or "
                  "is_rush_hour. The model can still learn rush-hour patterns implicitly (Gradient Boosting can "
                  "split on hour ranges), but no explicit derived binary flag exists in the feature set today. "
                  "This is the main gap versus the Week 1 task description, which is a good improvement to "
                  "mention if asked.", italic=True, color=RED)

    # ---------------- Pipeline construction ----------------
    add_heading(doc, "6. How the Pipeline Is Built (pipeline/train.py)")
    add_para(doc, "The training pipeline runs as a single script and performs these steps in order:")
    add_bullets(doc, [
        "1. Regenerate the raw dataset (generate_data.main()) and re-validate it (validate.validate()) so training always starts from a known-clean, versioned dataset",
        "2. Load the cleaned rows from data/processed/trips_clean.csv",
        "3. Split into train/test (80/20, seed=412 from params.yaml) using scikit-learn's train_test_split",
        "4. Build the shared preprocessor (ColumnTransformer: StandardScaler + OneHotEncoder)",
        "5. For each candidate model, wrap the preprocessor + regressor in an sklearn Pipeline and fit it on the training split only",
        "6. Evaluate both models on the held-out test split (MAE, RMSE, R2)",
        "7. Log every run's parameters, metrics, and tags to MLflow (experiment: 'delivery-eta')",
        "8. Select the model with the lowest test-set RMSE",
        "9. Persist the winning fitted Pipeline (preprocessor + model together) plus metadata to artifacts/model.joblib",
        "10. Write reports/model_comparison.md, reports/training_report.json, and experiments/runs.jsonl",
    ])
    add_para(doc, "Because the preprocessor and the regressor are fit together as one Pipeline object, there is "
                  "no risk of 'train/serve skew' - the API loads and reuses the exact same fitted object, so a "
                  "request's raw JSON payload goes through identical scaling/encoding as it did during training.")

    # ---------------- Model selection rationale ----------------
    add_heading(doc, "7. Why Gradient Boosting Was Chosen Over Linear Regression")
    add_table(doc, ["Model", "MAE (minutes)", "RMSE (minutes)", "R2", "Result"], [
        ["Linear Regression", "3.85", "4.82", "0.938", "Not selected"],
        ["Gradient Boosting", "3.62", "4.46", "0.947", "Selected (best RMSE)"],
    ], col_widths=[1.8, 1.4, 1.4, 1.0, 1.6])
    add_para(doc, "Selection rule: the pipeline always picks the model with the lowest RMSE on the held-out "
                  "test set (see train.py: selected = min(results, key=lambda x: (x['rmse'], -x['r2']))). This is "
                  "an automatic, metric-driven decision, not a manual/subjective choice.")
    add_para(doc, "Why Gradient Boosting wins here:")
    add_bullets(doc, [
        "Linear Regression assumes a purely additive, straight-line relationship between each feature and ETA - it cannot capture interactions such as 'high traffic AND rush hour together add much more delay than either alone'",
        "Gradient Boosting builds an ensemble of shallow decision trees (max_depth=3, 100 estimators, learning_rate=0.05) that can split on combinations of features, capturing non-linear effects like rush-hour x traffic x weather interactions",
        "The ETA generation formula itself combines rush-hour, traffic, and weather factors in a non-purely-linear way, so a tree-based model naturally fits this pattern better",
        "Gradient Boosting achieved lower error on every metric (MAE, RMSE) and slightly higher R2 on the same held-out test set, so it was selected automatically",
    ])
    add_para(doc, "Trade-off to mention: Linear Regression is more interpretable (coefficients directly show each "
                  "feature's effect) and cheaper to train/serve, while Gradient Boosting is a bit more of a "
                  "'black box' but more accurate here. Given ETA prediction cares more about accuracy than "
                  "explainability, the RMSE-based automatic selection is a reasonable choice.", italic=True, color=GRAY)

    # ---------------- API + feedback loop ----------------
    add_heading(doc, "8. How the API and Feedback Loop Work")
    add_heading(doc, "8.1 Prediction path (api/server.py)", level=2)
    add_bullets(doc, [
        "Client sends a POST /predict with trip details (coordinates, distance, hour, weekday, weather, traffic_level)",
        "Pydantic validates the payload (ranges, allowed weather/traffic values) - invalid input returns HTTP 422 with details, never reaches the model",
        "The pre-loaded artifacts/model.joblib Pipeline transforms the features and predicts the ETA",
        "The API generates a unique request_id (UUID), measures latency, and appends a JSON record (timestamp, request_id, input hash, predicted_eta_minutes, model_version, latency_ms) to logs/predictions.jsonl",
        "The API responds with predicted_eta_minutes, model_version, latency_ms, and request_id",
    ])
    add_heading(doc, "8.2 Feedback path (the loop that closes the system)", level=2)
    add_bullets(doc, [
        "Later, the real outcome of that trip becomes known (in production: from the completed trip; in this demo: client/feedback_client.py simulates it using the same physics-based duration formula as the data generator, plus noise)",
        "The client posts to POST /feedback with {request_id, actual_eta_minutes} - the API appends this to logs/feedback.jsonl",
        "pipeline/accuracy_report.py joins logs/predictions.jsonl and logs/feedback.jsonl by request_id, computing real error = actual - predicted for every matched pair",
        "It reports overall MAE/RMSE, plus MAE/RMSE split into an older half and a newer half of the matched pairs, so a rising RMSE in the newer window signals real accuracy degradation (accuracy_drift_triggered = true once newer-window RMSE crosses 8.0 minutes)",
        "This is the actual 'log predictions vs. actual times' requirement from Week 4 - distinct from PSI-based drift, which only looks at whether inputs look different, not whether the model's answers are still correct",
    ])
    add_para(doc, "Why this matters: PSI/TVD drift (feature drift) can flag a false alarm even when the model is "
                  "still accurate, and can miss a real problem if the input distribution looks unchanged but the "
                  "underlying world has shifted (e.g., speed limits changed). The feedback loop is the only way "
                  "to directly measure whether predictions are still correct.", italic=True, color=GRAY)

    # ---------------- Monitoring & observability walkthrough ----------------
    add_heading(doc, "9. Monitoring and Observability Dashboards - What Each Chart Shows")
    add_heading(doc, "9.1 Health / Monitoring Dashboard (reports/dashboard.html)", level=2)
    add_para(doc, "Built by pipeline/build_dashboard.py from logs/client_requests.jsonl (every request the "
                  "simulated client made, whether it succeeded, failed validation, or could not reach the API).")
    add_bullets(doc, [
        "Summary cards: Total requests (202), Uptime % (99.01%), 200 OK count (200), HTTP errors (1), Crashes (1)",
        "Latency-over-time line chart: each point is one request, colored by outcome (green = 200 OK, orange = HTTP error, red = crash/connection failure). A visible spike to ~4,000 ms in the middle of the timeline is the deliberately-injected crash test (the client tried to reach a wrong port and timed out) - everything else stays under ~10 ms",
        "Status code / crash timeline bar chart: green bars at height 200 for every successful call; the one orange bar (~422) marks the single injected invalid-payload test; a crash would show as a 0-height gap (not visible as a bar) at that same position",
    ])
    add_para(doc, "How to read it live: uptime near 100% and flat low latency = healthy system; any red points or "
                  "bar-height anomalies are what an on-call engineer would investigate first.", italic=True, color=GRAY)

    add_heading(doc, "9.2 Observability Dashboard (reports/observability_dashboard.html)", level=2)
    add_para(doc, "Built by pipeline/build_observability_dashboard.py from reports/drift_table.json (input "
                  "feature drift) and reports/accuracy_drift.json (predicted-vs-actual accuracy drift).")
    add_para(doc, "Part A - Input feature drift:", bold=True)
    add_bullets(doc, [
        "Summary cards: Concept drift PSI proxy (0.1221, green/ok), PSI threshold (0.2), TVD threshold (0.15), Numeric features drifting (0/7), Categorical features drifting (1/2)",
        "PSI per numeric feature bar chart: all 7 numeric features (pickup/dropoff lat-lon, distance, hour, weekday) sit well under the dashed 0.2 threshold line - no numeric drift detected with 200 live samples",
        "Total Variation Distance per categorical feature: weather is red/above threshold (0.257 > 0.15) because the simulated client over-samples 'clear' weather relative to the uniformly-distributed training data; traffic_level is green/ok (0.082)",
        "Numeric feature drift detail table and categorical feature drift detail table: the exact baseline vs. live mean/std/PSI/TVD numbers backing the charts above, with a color-coded ok/DRIFT status per row",
    ])
    add_para(doc, "Part B - Predicted vs. actual accuracy drift (the feedback loop, visualized):", bold=True)
    add_bullets(doc, [
        "Predicted vs. actual ETA line chart: blue line = model's predicted_eta_minutes, green line = simulated actual_eta_minutes for the same requests over time. When the two lines track closely, the model is accurate; visible gaps (actual consistently above predicted) show the model under-predicting, most pronounced when the surge simulation (--surge flag in feedback_client.py) is active",
        "Absolute prediction error bar chart: purple bars = |actual - predicted| in minutes per request, with a dashed line at the RMSE drift threshold (8.0 minutes). Many bars exceeding the dashed line indicate the model's real-world error has grown beyond the acceptable threshold - this is exactly what flips accuracy_drift_triggered to YES",
        "Recent predicted vs. actual detail table: the raw request_id / timestamp / predicted / actual / error / abs_error rows behind the charts, for the last 100 matched pairs",
    ])
    add_para(doc, "How to read it live: green PSI/TVD bars under the dashed line = inputs look like training data; "
                  "the predicted-vs-actual chart is the more important one operationally - if the two lines "
                  "diverge and the error bars regularly cross the threshold, that is a genuine signal to "
                  "retrain, independent of whether the PSI charts show anything unusual.", italic=True, color=GRAY)

    # ---------------- Demo script ----------------
    add_heading(doc, "10. Suggested Live Demo Script")
    add_table(doc, ["Step", "Command / Action", "What it Shows"], [
        ["1", "run.bat", "Full setup: build Docker image, start MLflow, train + select best model, run monitoring, run tests, start API"],
        ["2", "Open http://localhost:5000", "MLflow experiment tracking - all training runs, params, metrics"],
        ["3", "Open http://localhost:8000/docs", "Swagger UI - call /predict live with sample trip data"],
        ["4", "run_client.bat 150", "Simulate 150 real-time requests, submit feedback, compute drift + accuracy reports, rebuild both dashboards"],
        ["5", "reports/dashboard.html", "System health: uptime %, status-code timeline, latency, crash detection"],
        ["6", "reports/observability_dashboard.html", "PSI/TVD per feature (input drift) + predicted-vs-actual accuracy drift (feedback loop)"],
    ], col_widths=[0.6, 2.6, 3.3])

    # ---------------- Overall summary ----------------
    add_heading(doc, "11. Overall Coverage Summary")
    add_table(doc, ["Week", "Module", "Coverage"], [
        ["Week 1", "M2 - Ingest / Validate / Feature / Version", "~80% (feature engineering is basic, raw columns only)"],
        ["Week 2", "M3 - Train / Compare / Track Experiments", "100%"],
        ["Week 3", "M4 - Package / Serve REST API", "100%"],
        ["Week 4", "M5 - Log vs Actual / Drift / Monitor / Retrain Trigger", "~85% (retraining trigger is a signal, not yet automated)"],
    ], col_widths=[1.0, 3.5, 3.0])

    doc.save(OUT_PATH)
    print(f"Saved to {OUT_PATH}")


if __name__ == "__main__":
    build()
